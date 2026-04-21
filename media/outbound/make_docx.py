from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    # Parse **bold** markers
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    return p

# Title
title = doc.add_heading('AI培训复盘总结', level=0)
for run in title.runs:
    run.font.name = '微软雅黑'

# Subtitle / note
p = doc.add_paragraph()
run = p.add_run('复盘记录：给谁讲、讲什么、怎么讲 / 讲得怎么样、为什么这么讲')
run.italic = True
run.font.name = '微软雅黑'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============ 一、核心问题诊断 ============
add_heading('一、核心问题诊断', level=1)

add_heading('1. 准备阶段：培训前检查与知识储备不足', level=2)
add_bullet('**设备检查疏漏：** 误以为麦克风、摄像头已开启，实际未开启，导致培训开场信息传递失败，影响学员第一印象。')
add_bullet('**专业知识储备不扎实：** 对 AI 客服相关知识点掌握不够熟练，讲解时底气不足，难以应对学员提问。')
add_bullet('**培训重点未提前梳理：** 未明确哪些是核心内容（如快捷键等实操要点），导致讲解时主次不分。')
add_bullet('**学员背景了解不足：** 未针对新同事的认知水平提前设计讲解深度与节奏。')

add_heading('2. 执行阶段：节奏与专业度不足', level=2)
add_bullet('**进度过快：** 整体节奏拉得太快，学员来不及消化，重点部分一带而过。')
add_bullet('**讲解颗粒度粗：** 快捷键等实操细节没有细致展开，学员难以落地使用。')
add_bullet('**互动反馈缺失：** 未及时确认学员是否跟上，缺少答疑环节。')

add_heading('3. 后续跟进：闭环管理缺失', level=2)
add_bullet('**学习效果未验证：** 培训结束后没有通过考核或实操检验学员掌握情况。')
add_bullet('**待入职同事跟进不到位：** 对已培训但尚未入职的同事缺少持续跟进，容易出现信息断层或流失。')

# ============ 二、提升方向与行动方案 ============
add_heading('二、提升方向与行动方案', level=1)

add_heading('1. 前期准备流程', level=2)
add_bullet('**设备检查清单化：** 培训开始前 5 分钟完成"麦克风 / 摄像头 / 屏幕共享 / 网络"四项测试，并在群内发一句试音确认。')
add_bullet('**知识预习机制：** 培训前自己先把 AI 客服知识点过一遍——只有自己懂了，才能讲明白。对不熟的点提前查证，避免现场"卡壳"。')
add_bullet('**培训大纲与重点标记：** 提前列出培训目录，标注"重点 / 难点 / 必会"，让学员心里有地图。')
add_bullet('**学员信息同步：** 提前了解新同事的岗位、基础、入职时间，有针对性地调整讲解深度。')

add_heading('2. 塑造专业、清晰的讲解形象', level=2)
add_bullet('**节奏控制：**')
add_bullet('为每个模块设定时间节点（如：业务介绍 15 分钟 / 快捷键演示 20 分钟 / 答疑 10 分钟）。', level=1)
add_bullet('重点内容放慢，讲一遍 → 演示一遍 → 让学员跟做一遍。', level=1)
add_bullet('**快捷键等实操讲解标准化：**')
add_bullet('不只念快捷键，还要讲"什么场景下用 → 按哪几个键 → 实际效果"三件套。', level=1)
add_bullet('现场演示 + 截图/录屏留档，方便学员回看。', level=1)
add_bullet('**互动与反馈：**')
add_bullet('每讲完一个模块，问一句"这里还有疑问吗？"或让学员复述一遍。', level=1)
add_bullet('设计 1-2 个小练习或情景题，避免单向灌输。', level=1)

add_heading('3. 深入知识内核，做价值输出', level=2)
add_bullet('**AI 客服专业知识沉淀：**')
add_bullet('建立自己的"知识储备清单"，把常见问题、易错点、业务逻辑整理成文档，每次培训前刷一遍。', level=1)
add_bullet('遇到回答不上的问题，事后复盘并补充进知识库。', level=1)
add_bullet('**快捷键手册化：** 整理一份《AI 客服常用快捷键速查表》，培训后发给学员，方便日常查阅。')
add_bullet('**培训复用：** 每次培训后沉淀课件，下次直接迭代升级，避免重复造轮子。')

add_heading('4. 培训后跟进闭环', level=2)
add_bullet('**考核验收：** 培训结束安排简短测试或实操，确认学员掌握情况，薄弱点再单独辅导。')
add_bullet('**待入职同事跟进机制：**')
add_bullet('建立跟进表：姓名 / 入职状态 / 最近联系时间 / 下次跟进时间。', level=1)
add_bullet('每周固定时间 @ 一次，保持连接，降低流失风险。', level=1)
add_bullet('**培训效果回访：** 入职后 1 周、1 个月回访一次，听学员反馈哪些讲得不够清楚，反哺下一次培训。')

# Summary
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('核心一句话总结：')
run.bold = True
run.font.name = '微软雅黑'
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('培训的本质是让对方学会，不是讲完就行。���己先吃透、节奏要慢下来、重点要抠细节、后续要跟到底。')
run.font.name = '微软雅黑'
run.font.size = Pt(11)
run.italic = True

out = '/Users/weidian/openclaw-portable/config/workspace-bot1-agent/media/outbound/AI培训复盘总结.docx'
doc.save(out)
print(out)
